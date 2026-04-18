import sys
import os
import pandas as pd
from docx import Document
from datetime import datetime

from openpyxl import load_workbook
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.utils import get_column_letter

from PySide6.QtWidgets import (
    QApplication, QWidget, QLabel, QLineEdit, QTextEdit,
    QPushButton, QVBoxLayout, QHBoxLayout, QComboBox,
    QMessageBox, QTableWidget, QTableWidgetItem
)
from PySide6.QtCore import Qt

# ====== إعدادات ======

def load_reviewers():
    # جرب يقرأ من نفس فولدر exe الأول
    base_path = os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.getcwd()

    file_path = os.path.join(base_path, "reviewers_master.xlsx")

    if not os.path.exists(file_path):
        print("Reviewers file NOT FOUND:", file_path)
        return []

    df = pd.read_excel(file_path, dtype=str)

    if "FULL_NAME" not in df.columns:
        print("FULL_NAME column NOT FOUND")
        return []

    return df["FULL_NAME"].dropna().tolist()

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

TEMPLATE = resource_path("templates/template.docx")
OUTPUT_DIR = "output"

JOURNAL_FILES = {
    "مجلة السادات للبحوث الإدارية والمالية": os.path.join("databases", "السادات.xlsx"),
    "مجلة البحوث الإدارية": os.path.join("databases", "البحوث.xlsx"),
}

PREFIX_MAP = {
    "مجلة السادات للبحوث الإدارية والمالية": "JSA",
    "مجلة البحوث الإدارية": "JSO",
}

ISSUES = ["يناير", "أبريل", "يوليو", "أكتوبر"]

MONTHS = ["يناير","فبراير","مارس","أبريل","مايو","يونيو",
          "يوليو","أغسطس","سبتمبر","أكتوبر","نوفمبر","ديسمبر"]



os.makedirs(OUTPUT_DIR, exist_ok=True)

# ====== تنسيق Excel ======
def format_excel(file_path):
    wb = load_workbook(file_path)
    ws = wb.active

    last_row = ws.max_row
    last_col = ws.max_column
    table_ref = f"A1:{get_column_letter(last_col)}{last_row}"

    for t in list(ws.tables):
        del ws.tables[t]

    table = Table(displayName="DataTable", ref=table_ref)

    style = TableStyleInfo(
        name="TableStyleMedium2",
        showRowStripes=True,
        showColumnStripes=False
    )

    table.tableStyleInfo = style
    ws.add_table(table)
    ws.freeze_panes = "A2"

    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = max_length + 3

    wb.save(file_path)

# ====== وظائف ======

def count_issue_records(journal, issue_value):
    file_path = JOURNAL_FILES[journal]

    if not os.path.exists(file_path):
        return 0

    df = pd.read_excel(file_path, dtype=str)

    return len(df[df["ISSUE"] == issue_value])


def replace_placeholders(doc, mapping):
    for p in doc.paragraphs:
        for key, value in mapping.items():
            if key in p.text:
                for run in p.runs:
                    run.text = run.text.replace(key, value)

def generate_serial(journal):
    file_path = JOURNAL_FILES[journal]
    prefix = PREFIX_MAP[journal]

    if not os.path.exists(file_path):
        return f"{prefix}-26001"

    df = pd.read_excel(file_path)

    max_number = 0
    for serial in df["SERIAL"].astype(str):
        try:
            if serial.startswith(prefix):
                number = int(serial.split("-")[1])
                max_number = max(max_number, number)
        except:
            continue

    return f"{prefix}-{max_number + 1}"

def check_duplicate_name(journal, name_value):
    file_path = JOURNAL_FILES[journal]

    if not os.path.exists(file_path):
        return False

    df = pd.read_excel(file_path)
    name_value = name_value.strip().lower()

    return any(str(n).strip().lower() == name_value for n in df["NAME"])

def save_to_excel(journal, data):
    file_path = JOURNAL_FILES[journal]

    if os.path.exists(file_path):
        df = pd.read_excel(file_path, dtype=str)
    else:
        df = pd.DataFrame(columns=data.keys())

    df = pd.concat([df, pd.DataFrame([data])], ignore_index=True)
    df["ACCEPT_DATE"] = df["ACCEPT_DATE"].astype(str)

    df.to_excel(file_path, index=False)
    format_excel(file_path)

def generate_doc(data):
    template_path = os.path.join(os.getcwd(), "templates", "template.docx")

    if not os.path.exists(template_path):
        QMessageBox.critical(None, "خطأ", f"Template not found:\n{template_path}")
        return

    doc = Document(template_path)
    replace_placeholders(doc, data)

    filename = f"{data['{{SERIAL}}']} - {data['{{NAME}}']}.docx"
    doc.save(os.path.join(OUTPUT_DIR, filename))

# ====== شاشة التعديل ======

class EditForm(QWidget):
    def __init__(self, record, file_path):
        super().__init__()

        self.record = record
        self.file_path = file_path

        self.setWindowTitle("تعديل")
        self.resize(500, 500)
        self.setLayoutDirection(Qt.RightToLeft)

        layout = QVBoxLayout()

        # ===== اسم الباحث =====
        self.name = QLineEdit(record["NAME"])

        layout.addWidget(QLabel("اسم الباحث"))
        layout.addWidget(self.name)

        # ===== عنوان البحث =====
        self.title = QTextEdit(record["TITLE"])

        layout.addWidget(QLabel("عنوان البحث"))
        layout.addWidget(self.title)

        # ===== التاريخ =====
        self.day = QComboBox()
        self.day.addItems([str(i) for i in range(1, 32)])

        self.month = QComboBox()
        self.month.addItems(MONTHS)

        self.year = QComboBox()
        self.year.addItems([str(y) for y in range(2025, 2031)])

        date_layout = QHBoxLayout()
        date_layout.addWidget(self.day)
        date_layout.addWidget(self.month)
        date_layout.addWidget(self.year)

        layout.addWidget(QLabel("تاريخ القبول"))
        layout.addLayout(date_layout)

        # تحميل التاريخ القديم
        try:
            d, m, y = record["ACCEPT_DATE"].split("/")
            self.day.setCurrentText(str(int(d)))
            self.month.setCurrentText(MONTHS[int(m) - 1])
            self.year.setCurrentText(y)
        except:
            pass

        # ===== العدد =====
        self.issue = QComboBox()
        self.issue.addItems(ISSUES)

        self.issue_year = QComboBox()
        self.issue_year.addItems([str(y) for y in range(2025, 2031)])

        issue_layout = QHBoxLayout()
        issue_layout.addWidget(self.issue)
        issue_layout.addWidget(self.issue_year)

        layout.addWidget(QLabel("العدد"))
        layout.addLayout(issue_layout)

        # تحميل العدد القديم
        try:
            issue_name, issue_year = record["ISSUE"].split()
            self.issue.setCurrentText(issue_name)
            self.issue_year.setCurrentText(issue_year)
        except:
            pass
        
        # ===== تحميل المحكمين =====
        
        reviewers_list = load_reviewers()

        self.reviewer1 = QComboBox()
        self.reviewer1.addItem("")  # قيمة فاضية
        self.reviewer1.addItems(reviewers_list)

        reviewer1 = record.get("REVIEWER1", "")
        if pd.isna(reviewer1):
            reviewer1 = ""

        self.reviewer1.setCurrentText(str(reviewer1))


        self.reviewer2 = QComboBox()
        self.reviewer2.addItem("")
        self.reviewer2.addItems(reviewers_list)

        reviewer2 = record.get("REVIEWER2", "")
        if pd.isna(reviewer2):
            reviewer2 = ""

        self.reviewer2.setCurrentText(str(reviewer2))


        layout.addWidget(QLabel("المحكم الأول"))
        layout.addWidget(self.reviewer1)

        layout.addWidget(QLabel("المحكم الثاني"))
        layout.addWidget(self.reviewer2)

        # ===== زر الحفظ =====
        save_btn = QPushButton("حفظ + إعادة إصدار")
        save_btn.clicked.connect(self.save)

        layout.addWidget(save_btn)

        self.setLayout(layout)

    def save(self):
        name = self.name.text().strip()
        title = self.title.toPlainText().strip()

        if not name or not title:
            QMessageBox.warning(self, "خطأ", "اكمل البيانات")
            return

        day = self.day.currentText().zfill(2)
        month_index = MONTHS.index(self.month.currentText()) + 1
        month = str(month_index).zfill(2)
        year = self.year.currentText()

        accept_date = f"{day}/{month}/{year}"
        issue_full = f"{self.issue.currentText()} {self.issue_year.currentText()}"

        df = pd.read_excel(self.file_path, dtype=str)

        serial = self.record["SERIAL"]

        df.loc[df["SERIAL"] == serial, "NAME"] = name
        df.loc[df["SERIAL"] == serial, "TITLE"] = title
        df.loc[df["SERIAL"] == serial, "ACCEPT_DATE"] = accept_date
        df.loc[df["SERIAL"] == serial, "ISSUE"] = issue_full
        df.loc[df["SERIAL"] == serial, "REVIEWER1"] = self.reviewer1.currentText()
        df.loc[df["SERIAL"] == serial, "REVIEWER2"] = self.reviewer2.currentText()

        df.to_excel(self.file_path, index=False)
        format_excel(self.file_path)

        doc_data = {
            "{{SERIAL}}": serial,
            "{{NAME}}": name,
            "{{JOURNAL}}": self.record["JOURNAL"],
            "{{TITLE}}": title,
            "{{ACCEPT_DATE}}": accept_date,
            "{{ISSUE}}": issue_full,
        }

        generate_doc(doc_data)

        QMessageBox.information(self, "تم", "تم التعديل وإعادة إصدار الخطاب")
        self.close()
        
# ====== شاشة البحث ======

class SearchWindow(QWidget):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("بحث / تعديل")
        self.resize(900, 600)
        self.setLayoutDirection(Qt.RightToLeft)

        layout = QVBoxLayout()

        # ===== اختيار المجلة =====
        self.journal = QComboBox()
        self.journal.addItems(JOURNAL_FILES.keys())

        layout.addWidget(QLabel("المجلة"))
        layout.addWidget(self.journal)

        # ===== البحث =====
        self.search_input = QLineEdit()

        search_btn = QPushButton("بحث")
        search_btn.clicked.connect(self.search)

        layout.addWidget(QLabel("اسم الباحث"))
        layout.addWidget(self.search_input)
        layout.addWidget(search_btn)

        # ===== الجدول =====
        self.table = QTableWidget()
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        self.table.setAlternatingRowColors(True)
        layout.addWidget(self.table)

        # ===== الأزرار =====
        edit_btn = QPushButton("تعديل")
        edit_btn.clicked.connect(self.edit_record)

        delete_btn = QPushButton("حذف")
        delete_btn.clicked.connect(self.delete_record)

        layout.addWidget(edit_btn)
        layout.addWidget(delete_btn)

        self.setLayout(layout)

    def search(self):
        journal = self.journal.currentText()
        file_path = JOURNAL_FILES[journal]

        if not os.path.exists(file_path):
            QMessageBox.warning(self, "خطأ", "لا يوجد بيانات")
            return

        df = pd.read_excel(file_path, dtype=str)

        keyword = self.search_input.text().strip().lower()

        if keyword:
            df = df[df["NAME"].str.lower().str.contains(keyword)]

        self.data = df
        self.file_path = file_path

        self.table.setRowCount(len(df))
        self.table.setColumnCount(len(df.columns))
        self.table.setHorizontalHeaderLabels(df.columns)

        for i in range(len(df)):
            for j in range(len(df.columns)):
                self.table.setItem(i, j, QTableWidgetItem(str(df.iloc[i, j])))

    def edit_record(self):
        row = self.table.currentRow()

        if row == -1:
            QMessageBox.warning(self, "خطأ", "اختار صف")
            return

        record = self.data.iloc[row]
        self.edit_form = EditForm(record, self.file_path)
        self.edit_form.show()

    def delete_record(self):
        row = self.table.currentRow()

        if row == -1:
            QMessageBox.warning(self, "خطأ", "اختار صف")
            return

        if QMessageBox.question(self, "تأكيد", "هل أنت متأكد؟") != QMessageBox.Yes:
            return

        serial = self.table.item(row, 0).text()

        df = pd.read_excel(self.file_path, dtype=str)
        df = df[df["SERIAL"] != serial]

        df.to_excel(self.file_path, index=False)
        format_excel(self.file_path)

        QMessageBox.information(self, "تم", "تم الحذف")
        self.search()

# ====== الشاشة الرئيسية ======

class App(QWidget):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("نظام خطابات القبول")
        self.resize(900, 700)
        self.setMinimumSize(800, 600)
        self.setLayoutDirection(Qt.RightToLeft)

        layout = QVBoxLayout()

        # ===== المجلة =====
        self.journal = QComboBox()
        self.journal.addItems(JOURNAL_FILES.keys())

        layout.addWidget(QLabel("المجلة"))
        layout.addWidget(self.journal)

        # ===== اسم الباحث =====
        self.name = QLineEdit()

        layout.addWidget(QLabel("اسم الباحث"))
        layout.addWidget(self.name)

        # ===== عنوان البحث =====
        self.title = QTextEdit()

        layout.addWidget(QLabel("عنوان البحث"))
        layout.addWidget(self.title)

        # ===== التاريخ =====
        self.day = QComboBox()
        self.day.addItems([str(i) for i in range(1, 32)])

        self.month = QComboBox()
        self.month.addItems(MONTHS)

        self.year = QComboBox()
        self.year.addItems([str(y) for y in range(2025, 2031)])

        date_layout = QHBoxLayout()
        date_layout.addWidget(self.day)
        date_layout.addWidget(self.month)
        date_layout.addWidget(self.year)

        today_btn = QPushButton("اليوم")
        today_btn.clicked.connect(self.set_today)

        layout.addWidget(QLabel("تاريخ القبول"))
        layout.addLayout(date_layout)
        layout.addWidget(today_btn)

        # ===== العدد =====
        self.issue = QComboBox()
        self.issue.addItems(ISSUES)

        self.issue_year = QComboBox()
        self.issue_year.addItems([str(y) for y in range(2025, 2031)])

        issue_layout = QHBoxLayout()
        issue_layout.addWidget(self.issue)
        issue_layout.addWidget(self.issue_year)

        layout.addWidget(QLabel("العدد"))
        layout.addLayout(issue_layout)
        
        self.issue_count_label = QLabel("عدد الأبحاث: 0 / 35")
        layout.addWidget(self.issue_count_label)
        
        
       # ===== المحكمين =====
        reviewers_list = load_reviewers()

        self.reviewer1 = QComboBox()
        self.reviewer1.addItem("")  # قيمة فاضية
        self.reviewer1.addItems(reviewers_list)

        self.reviewer2 = QComboBox()
        self.reviewer2.addItem("")
        self.reviewer2.addItems(reviewers_list)

        layout.addWidget(QLabel("المحكم الأول"))
        layout.addWidget(self.reviewer1)

        layout.addWidget(QLabel("المحكم الثاني"))
        layout.addWidget(self.reviewer2)
        

        # ===== الأزرار =====
        btn_generate = QPushButton("إصدار")
        btn_generate.clicked.connect(self.submit)

        btn_search = QPushButton("بحث / تعديل")
        btn_search.clicked.connect(self.open_search)

        layout.addWidget(btn_generate)
        layout.addWidget(btn_search)
        
        self.issue.currentIndexChanged.connect(self.update_issue_count)
        self.issue_year.currentIndexChanged.connect(self.update_issue_count)
        self.journal.currentIndexChanged.connect(self.update_issue_count)
        
        self.update_issue_count()

        self.setLayout(layout)

    def set_today(self):
        today = datetime.now()
        self.day.setCurrentText(str(today.day))
        self.month.setCurrentText(MONTHS[today.month - 1])
        self.year.setCurrentText(str(today.year))
        
    def update_issue_count(self):
        journal = self.journal.currentText()
        issue_full = f"{self.issue.currentText()} {self.issue_year.currentText()}"

        count = count_issue_records(journal, issue_full)

        self.issue_count_label.setText(f"عدد الأبحاث في هذا العدد: {count} / 35")

        # تحسين اللون
        if count >= 35:
            self.issue_count_label.setStyleSheet("color: red; font-weight: bold;")
        else:
            self.issue_count_label.setStyleSheet("color: green;")

    def submit(self):
        journal = self.journal.currentText()
        name = self.name.text().strip()
        title = self.title.toPlainText().strip()

        if not name or not title:
            QMessageBox.warning(self, "خطأ", "اكمل البيانات")
            return

        if check_duplicate_name(journal, name):
            if QMessageBox.question(self, "تحذير", "الاسم موجود قبل كده، تكمل؟") != QMessageBox.Yes:
                return

        serial = generate_serial(journal)

        day = self.day.currentText().zfill(2)
        month_index = MONTHS.index(self.month.currentText()) + 1
        month = str(month_index).zfill(2)
        year = self.year.currentText()

        accept_date = f"{day}/{month}/{year}"
        issue_full = f"{self.issue.currentText()} {self.issue_year.currentText()}"
        
        count = count_issue_records(journal, issue_full)

        if count >= 35:
            QMessageBox.warning(
                self,
                "تنبيه",
                f"تم الوصول للحد الأقصى (35 بحث)\nعدد الأبحاث الحالي: {count}"
            )
            return

        excel_data = {
            "SERIAL": serial,
            "NAME": name,
            "JOURNAL": journal,
            "TITLE": title,
            "ACCEPT_DATE": accept_date,
            "ISSUE": issue_full,
            "REVIEWER1": self.reviewer1.currentText(),
            "REVIEWER2": self.reviewer2.currentText(),
        }
        doc_data = {
            "{{SERIAL}}": serial,
            "{{NAME}}": name,
            "{{JOURNAL}}": journal,
            "{{TITLE}}": title,
            "{{ACCEPT_DATE}}": accept_date,
            "{{ISSUE}}": issue_full,
        }

        save_to_excel(journal, excel_data)
        generate_doc(doc_data)
        self.update_issue_count()

        QMessageBox.information(self, "تم", f"تم إنشاء الخطاب برقم {serial}")

        self.name.clear()
        self.title.clear()
        self.reviewer1.clear()
        self.reviewer2.clear()

    def open_search(self):
        self.search_window = SearchWindow()
        self.search_window.show()

# ====== تشغيل ======

if __name__ == "__main__":
    app = QApplication(sys.argv)

    # 🔥 ستايل موحد
    app.setStyleSheet("""
    QWidget { font-size: 14pt; }
    QPushButton { height: 40px; font-weight: bold; }
    QLineEdit, QTextEdit, QComboBox { height: 35px; }
    """)

    window = App()
    window.show()
    sys.exit(app.exec())